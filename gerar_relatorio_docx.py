from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

figures = [
    (1, r"figuras/fig01_index.png", "Página inicial do dashboard (index.html)"),
    (2, r"figuras/fig02_total_acidentes_ano.png", "Total de acidentes por ano (2023–2025)"),
    (3, r"figuras/fig03_vitimas_gravidade.png", "Distribuição de vítimas por gravidade"),
    (4, r"figuras/fig04_rodovias_acidentes.png", "Top rodovias com mais acidentes"),
    (5, r"figuras/fig05_heatmap_rodovia_mes.png", "Heatmap rodovia × mês"),
    (6, r"figuras/fig06_matriz_correlacao.png", "Matriz de correlação entre variáveis"),
    (7, r"figuras/fig07_taxa_mortalidade.png", "Taxa de mortalidade por rodovia"),
]

repo = r"C:\vpcapanema\acidentes-der-analise-exploratoria"

# Create document
_doc = Document()

# Title
p = _doc.add_paragraph()
run = p.add_run("RELATÓRIO DE ATIVIDADES MENSAL\n")
run.bold = True
run.font.size = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p2 = _doc.add_paragraph()
run2 = p2.add_run("Projeto: Análise de Acidentes DER-SP\n")
run2.font.size = Pt(12)
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

p3 = _doc.add_paragraph()
run3 = p3.add_run("Período: 11/12/2025 a 11/01/2026")
run3.font.size = Pt(11)
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadata
_doc.add_paragraph("")
meta = _doc.add_paragraph()
meta.add_run("Responsável: Vinicius Capanema\n").bold = True
meta.add_run("Status: Concluído\n").bold = True
meta.add_run("Versão: 1.0\n").bold = True

# Lista de Figuras
_doc.add_heading("Lista de Figuras", level=1)
for num, path, caption in figures:
    _doc.add_paragraph(f"Figura {num} - {caption}")

# Resumo executivo
_doc.add_heading("1. Resumo Executivo", level=1)
_doc.add_paragraph(
    "No período de 11/12/2025 a 11/01/2026 foi desenvolvida a plataforma "
    "'Análise de Acidentes DER-SP', com consolidação de dados de 2023 a 2025 "
    "e geração de 29 gráficos interativos. A aplicação permite análise temporal, "
    "geográfica e de severidade, oferecendo subsídios para decisões baseadas em dados. "
    "A visão geral da aplicação é apresentada na Figura 1."
)

# Aplicação e visão geral
_doc.add_heading("2. Aplicação e Visão Geral", level=1)
_doc.add_paragraph(
    "A página inicial reúne o acesso central aos dashboards e relatórios. "
    "A Figura 1 destaca a interface principal com navegação por seções e acesso rápido "
    "a gráficos temáticos."
)

# Figura 1
img_path = f"{repo}\\{figures[0][1]}"
_doc.add_picture(img_path, width=Inches(6.3))
cap = _doc.add_paragraph(f"Figura 1 - {figures[0][2]}")
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Análises principais
_doc.add_heading("3. Principais Análises", level=1)
_doc.add_paragraph(
    "A série histórica de acidentes evidencia a evolução anual dos registros, "
    "conforme apresentado na Figura 2. Em complemento, a distribuição de vítimas "
    "por gravidade (Figura 3) orienta ações de prevenção e alocação de recursos."
)

# Figura 2
img_path = f"{repo}\\{figures[1][1]}"
_doc.add_picture(img_path, width=Inches(6.3))
cap = _doc.add_paragraph(f"Figura 2 - {figures[1][2]}")
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Figura 3
img_path = f"{repo}\\{figures[2][1]}"
_doc.add_picture(img_path, width=Inches(6.3))
cap = _doc.add_paragraph(f"Figura 3 - {figures[2][2]}")
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Rodovias críticas
_doc.add_paragraph(
    "A identificação de rodovias críticas permite priorizar investimentos. "
    "A Figura 4 mostra o ranking das rodovias com maior número de ocorrências."
)

# Figura 4
img_path = f"{repo}\\{figures[3][1]}"
_doc.add_picture(img_path, width=Inches(6.3))
cap = _doc.add_paragraph(f"Figura 4 - {figures[3][2]}")
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Heatmap e correlação
_doc.add_paragraph(
    "Para investigar padrões sazonais e relações entre variáveis, "
    "foram utilizados heatmaps e matrizes de correlação. "
    "A Figura 5 apresenta o heatmap rodovia × mês e a Figura 6 mostra a matriz de correlação."
)

# Figura 5
img_path = f"{repo}\\{figures[4][1]}"
_doc.add_picture(img_path, width=Inches(6.3))
cap = _doc.add_paragraph(f"Figura 5 - {figures[4][2]}")
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Figura 6
img_path = f"{repo}\\{figures[5][1]}"
_doc.add_picture(img_path, width=Inches(6.3))
cap = _doc.add_paragraph(f"Figura 6 - {figures[5][2]}")
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Mortalidade
_doc.add_paragraph(
    "A análise de risco é sintetizada pela taxa de mortalidade por rodovia, "
    "evidenciando vias prioritárias para ações de segurança (Figura 7)."
)

# Figura 7
img_path = f"{repo}\\{figures[6][1]}"
_doc.add_picture(img_path, width=Inches(6.3))
cap = _doc.add_paragraph(f"Figura 7 - {figures[6][2]}")
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metodologia
_doc.add_heading("4. Metodologia", level=1)
_doc.add_paragraph(
    "O processo envolveu: (i) consolidação das bases anuais, (ii) limpeza e padronização, "
    "(iii) análise exploratória e (iv) geração de 29 gráficos interativos com Plotly. "
    "Os scripts Python foram modularizados para facilitar manutenção e expansão."
)

# Entregáveis
_doc.add_heading("5. Entregáveis", level=1)
_doc.add_paragraph(
    "• 29 gráficos interativos em HTML\n"
    "• 3 portais web (index, portal, dashboard)\n"
    "• Dataset consolidado com 39.578 registros\n"
    "• Documentação e relatórios HTML\n"
)

# Conclusão
_doc.add_heading("6. Conclusão", level=1)
_doc.add_paragraph(
    "O projeto foi concluído com sucesso no período estabelecido, entregando uma plataforma "
    "robusta e acessível para análise de acidentes do DER-SP. As Figuras 1 a 7 comprovam a "
    "amplitude das análises e a qualidade das visualizações geradas."
)

# Save
out_path = f"{repo}\\RELATORIO_ATIVIDADES_11-12_11-01.docx"
_doc.save(out_path)
print(out_path)
